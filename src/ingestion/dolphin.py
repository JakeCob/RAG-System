"""Layout-aware document parser with multi-format support.

Handles: TXT, MD, PDF, DOCX, PPTX, XLSX, CSV
Preserves: Tables, headers, lists, images, page structure
"""

from __future__ import annotations

import io
import uuid
import zipfile
from pathlib import Path
from typing import Any, ClassVar, Literal, TypedDict
from xml.etree import ElementTree

import pandas as pd  # type: ignore[import-untyped]

from app.schemas import AgentFailure, ErrorCodes
from app.schemas.parser import ParsedChunk
from ingestion.base import BaseParser


partition_docx: Any | None = None
partition_pdf: Any | None = None
partition_pptx: Any | None = None

try:  # Optional unstructured dependency
    from unstructured.partition.docx import partition_docx as _partition_docx
    from unstructured.partition.pdf import partition_pdf as _partition_pdf
    from unstructured.partition.pptx import partition_pptx as _partition_pptx

    partition_docx = _partition_docx
    partition_pdf = _partition_pdf
    partition_pptx = _partition_pptx
except Exception:  # pragma: no cover - fallback paths handle missing deps
    partition_docx = None
    partition_pdf = None
    partition_pptx = None

pdfplumber: Any | None

try:  # Optional PDF fallback
    import pdfplumber
except Exception:  # pragma: no cover - fallback paths handle missing deps
    pdfplumber = None


ChunkType = Literal["text", "table"]
LayoutType = Literal["text", "table", "image", "header", "list"]


class ChunkData(TypedDict):
    text: str
    type: ChunkType


class DolphinParser(BaseParser):
    """Multi-format document parser with layout preservation."""

    SUPPORTED_EXTENSIONS: ClassVar[set[str]] = {
        ".txt",
        ".md",
        ".markdown",
        ".pdf",
        ".docx",
        ".pptx",
        ".xlsx",
        ".xls",
        ".csv",
    }

    def __init__(
        self,
        *,
        enable_ocr: bool = False,
        chunk_size: int = 1000,
    ) -> None:
        super().__init__()
        self.enable_ocr = enable_ocr
        self.chunk_size = chunk_size

    def _split_by_markdown_tables(self, text: str) -> list[ChunkData]:
        """Split Markdown into chunks, isolating contiguous table blocks.

        Args:
            text: Markdown/text content.

        Returns:
            List of chunk dicts with keys: `text` and `type` ("text" or "table").
        """
        lines = text.split("\n")
        chunks_data: list[ChunkData] = []
        current_lines: list[str] = []
        in_table = False

        for line in lines:
            stripped = line.strip()
            is_table_row = stripped.startswith("|") and stripped.endswith("|")

            if is_table_row:
                if not in_table:
                    if current_lines:
                        chunks_data.append(
                            {"text": "\n".join(current_lines), "type": "text"}
                        )
                        current_lines = []
                    in_table = True
                current_lines.append(line)
                continue

            if in_table:
                if current_lines:
                    chunks_data.append(
                        {"text": "\n".join(current_lines), "type": "table"}
                    )
                    current_lines = []
                in_table = False

            current_lines.append(line)

        if current_lines:
            chunks_data.append(
                {
                    "text": "\n".join(current_lines),
                    "type": "table" if in_table else "text",
                }
            )

        return chunks_data

    def parse(  # noqa: C901
        self, content: str | bytes, metadata: dict[str, Any]
    ) -> list[ParsedChunk] | AgentFailure:
        """Parse content, preserving tables as distinct chunks.

        Args:
            content: Raw file content (str or bytes).
            metadata: Metadata associated with the content (e.g., URL, title).

        Returns:
            A list of structured chunks or AgentFailure on error.
        """
        if not isinstance(metadata, dict):
            return self._failure(
                ErrorCodes.PARSER_INVALID_INPUT, "metadata must be a dict."
            )

        filename = str(metadata.get("filename", ""))
        file_ext = Path(filename).suffix.lower()
        if not file_ext:
            return self._parse_text(content, metadata)

        if file_ext in {".doc", ".ppt"}:
            return self._failure(
                ErrorCodes.PARSER_UNSUPPORTED_FORMAT,
                (
                    f"Legacy format {file_ext} not supported. "
                    f"Please convert to {file_ext}x."
                ),
            )

        if file_ext in {".html", ".htm"}:
            return self._failure(
                ErrorCodes.PARSER_UNSUPPORTED_FORMAT,
                "HTML files should be ingested via the Web Connector.",
            )

        if file_ext not in self.SUPPORTED_EXTENSIONS:
            return self._failure(
                ErrorCodes.PARSER_UNSUPPORTED_FORMAT,
                f"Unsupported file format: {file_ext}",
            )

        try:
            if file_ext in {".txt", ".md", ".markdown"}:
                return self._parse_text(content, metadata)
            if file_ext == ".pdf":
                return self._parse_pdf(content, metadata)
            if file_ext == ".docx":
                return self._parse_docx(content, metadata)
            if file_ext == ".pptx":
                return self._parse_pptx(content, metadata)
            if file_ext in {".xlsx", ".xls", ".csv"}:
                return self._parse_spreadsheet(content, metadata, file_ext)
        except Exception as exc:  # Defensive catch-all
            return self._failure(
                ErrorCodes.PARSER_INVALID_INPUT, f"Parsing failed: {exc}"
            )

        return self._failure(
            ErrorCodes.PARSER_UNSUPPORTED_FORMAT,
            f"Unsupported file format: {file_ext}",
        )

    def _parse_text(
        self, content: str | bytes, _metadata: dict[str, Any]
    ) -> list[ParsedChunk]:
        text = self._decode_content(content)
        if not text:
            raise ValueError("Content cannot be empty.")

        chunks_data = self._split_by_markdown_tables(text)
        parsed_chunks: list[ParsedChunk] = []
        chunk_index = 0

        for data in chunks_data:
            chunk_text = data["text"].strip()
            if not chunk_text:
                continue

            if data["type"] == "text":
                for split_text in self.chunk(chunk_text, limit=self.chunk_size):
                    if not split_text.strip():
                        continue
                    parsed_chunks.append(
                        ParsedChunk(
                            chunk_id=str(uuid.uuid4()),
                            content=split_text,
                            chunk_index=chunk_index,
                            layout_type="text",
                            bbox=None,
                        )
                    )
                    chunk_index += 1
            else:
                parsed_chunks.append(
                    ParsedChunk(
                        chunk_id=str(uuid.uuid4()),
                        content=chunk_text,
                        chunk_index=chunk_index,
                        layout_type="table",
                        bbox=None,
                    )
                )
                chunk_index += 1

        return parsed_chunks

    def _parse_pdf(
        self, content: str | bytes, _metadata: dict[str, Any]
    ) -> list[ParsedChunk] | AgentFailure:
        if not isinstance(content, bytes | bytearray):
            return self._failure(
                ErrorCodes.PARSER_INVALID_INPUT, "PDF content must be bytes."
            )

        if partition_pdf is not None:
            try:
                elements = partition_pdf(
                    file=io.BytesIO(content),
                    strategy="hi_res",
                    infer_table_structure=True,
                    extract_images_in_pdf=True,
                    ocr_languages="eng" if self.enable_ocr else None,
                )
                chunks = self._elements_to_chunks(elements)
                if self.enable_ocr and not chunks:
                    return self._failure(
                        ErrorCodes.PARSER_OCR_FAILED,
                        "OCR produced no text for PDF.",
                    )
                return chunks
            except Exception as exc:
                if self.enable_ocr and self._is_ocr_error(exc):
                    return self._failure(
                        ErrorCodes.PARSER_OCR_FAILED,
                        f"OCR failed for PDF: {exc}",
                    )
                if self._is_corrupted_pdf_error(exc):
                    return self._failure(
                        ErrorCodes.PARSER_CORRUPTED_FILE,
                        f"Failed to parse PDF: {exc}",
                    )
                return self._failure(
                    ErrorCodes.PARSER_INVALID_INPUT, f"Failed to parse PDF: {exc}"
                )

        return self._parse_pdf_fallback(content)

    def _parse_pdf_fallback(  # noqa: C901
        self, content: bytes
    ) -> list[ParsedChunk] | AgentFailure:
        if pdfplumber is None:
            return self._failure(
                ErrorCodes.PARSER_INVALID_INPUT,
                "PDF parsing unavailable: pdfplumber not installed.",
            )

        parsed_chunks: list[ParsedChunk] = []
        chunk_index = 0
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page_number, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text() or ""
                    tables = page.extract_tables() or []
                    for table in tables:
                        if not table:
                            continue
                        header = table[0]
                        rows = table[1:] if len(table) > 1 else []
                        df = pd.DataFrame(rows, columns=header)
                        markdown = self._table_to_markdown(df)
                        if markdown:
                            parsed_chunks.append(
                                ParsedChunk(
                                    chunk_id=str(uuid.uuid4()),
                                    content=markdown,
                                    chunk_index=chunk_index,
                                    layout_type="table",
                                    page_number=page_number,
                                    bbox=None,
                                )
                            )
                            chunk_index += 1

                    if page_text.strip():
                        for chunk in self._split_by_markdown_tables(page_text):
                            chunk_text = chunk["text"].strip()
                            if not chunk_text:
                                continue
                            layout: ChunkType = (
                                "table" if chunk["type"] == "table" else "text"
                            )
                            parsed_chunks.append(
                                ParsedChunk(
                                    chunk_id=str(uuid.uuid4()),
                                    content=chunk_text,
                                    chunk_index=chunk_index,
                                    layout_type=layout,
                                    page_number=page_number,
                                    bbox=None,
                                )
                            )
                            chunk_index += 1
        except Exception as exc:
            if self._is_corrupted_pdf_error(exc):
                return self._failure(
                    ErrorCodes.PARSER_CORRUPTED_FILE, f"Failed to parse PDF: {exc}"
                )
            return self._failure(
                ErrorCodes.PARSER_INVALID_INPUT, f"Failed to parse PDF: {exc}"
            )

        if self.enable_ocr and not parsed_chunks:
            return self._failure(
                ErrorCodes.PARSER_OCR_FAILED, "OCR produced no text for PDF."
            )

        return parsed_chunks

    def _parse_docx(  # noqa: C901
        self, content: str | bytes, _metadata: dict[str, Any]
    ) -> list[ParsedChunk] | AgentFailure:
        if not isinstance(content, bytes | bytearray):
            return self._failure(
                ErrorCodes.PARSER_INVALID_INPUT, "DOCX content must be bytes."
            )

        if partition_docx is not None:
            try:
                elements = partition_docx(file=io.BytesIO(content))
                return self._elements_to_chunks(elements)
            except Exception as exc:
                return self._failure(
                    ErrorCodes.PARSER_INVALID_INPUT, f"Failed to parse DOCX: {exc}"
                )

        try:
            from docx import Document
        except Exception as exc:  # pragma: no cover - dependency missing
            return self._failure(
                ErrorCodes.PARSER_INVALID_INPUT, f"DOCX parsing unavailable: {exc}"
            )

        try:
            doc = Document(io.BytesIO(content))
        except Exception as exc:
            return self._failure(
                ErrorCodes.PARSER_CORRUPTED_FILE, f"Failed to parse DOCX: {exc}"
            )

        parsed_chunks: list[ParsedChunk] = []
        chunk_index = 0

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name.startswith("Heading"):
                layout: LayoutType = "header"
            elif "List" in style_name:
                layout = "list"
            else:
                layout = "text"

            for chunk in self._split_by_markdown_tables(text):
                chunk_text = chunk["text"].strip()
                if not chunk_text:
                    continue
                chunk_layout: LayoutType = (
                    "table" if chunk["type"] == "table" else layout
                )
                parsed_chunks.append(
                    ParsedChunk(
                        chunk_id=str(uuid.uuid4()),
                        content=chunk_text,
                        chunk_index=chunk_index,
                        layout_type=chunk_layout,
                        bbox=None,
                    )
                )
                chunk_index += 1

        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header = rows[0]
            data_rows = rows[1:] if len(rows) > 1 else []
            df = pd.DataFrame(data_rows, columns=header)
            markdown = self._table_to_markdown(df)
            if not markdown:
                continue
            parsed_chunks.append(
                ParsedChunk(
                    chunk_id=str(uuid.uuid4()),
                    content=markdown,
                    chunk_index=chunk_index,
                    layout_type="table",
                    bbox=None,
                )
            )
            chunk_index += 1

        return parsed_chunks

    def _parse_pptx(
        self, content: str | bytes, _metadata: dict[str, Any]
    ) -> list[ParsedChunk] | AgentFailure:
        if not isinstance(content, bytes | bytearray):
            return self._failure(
                ErrorCodes.PARSER_INVALID_INPUT, "PPTX content must be bytes."
            )

        if partition_pptx is not None:
            try:
                elements = partition_pptx(file=io.BytesIO(content))
                return self._elements_to_chunks(elements)
            except Exception:
                pass

        return self._parse_pptx_zip(content)

    def _parse_pptx_zip(self, content: bytes) -> list[ParsedChunk] | AgentFailure:
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zipf:
                slide_paths = [
                    name
                    for name in zipf.namelist()
                    if name.startswith("ppt/slides/slide") and name.endswith(".xml")
                ]
                slide_paths.sort(key=self._slide_sort_key)

                parsed_chunks: list[ParsedChunk] = []
                chunk_index = 0
                for slide_path in slide_paths:
                    slide_number = self._slide_number_from_path(slide_path)
                    slide_text = self._extract_text_from_xml(
                        zipf.read(slide_path)
                    ).strip()

                    notes_path = (
                        f"ppt/notesSlides/notesSlide{slide_number}.xml"
                        if slide_number is not None
                        else None
                    )
                    notes_text = ""
                    if notes_path and notes_path in zipf.namelist():
                        notes_text = self._extract_text_from_xml(
                            zipf.read(notes_path)
                        ).strip()

                    content_parts = [part for part in [slide_text, notes_text] if part]
                    if not content_parts:
                        continue

                    combined = "\n\n".join(content_parts)
                    parsed_chunks.append(
                        ParsedChunk(
                            chunk_id=str(uuid.uuid4()),
                            content=combined,
                            chunk_index=chunk_index,
                            layout_type="text",
                            page_number=slide_number,
                            bbox=None,
                        )
                    )
                    chunk_index += 1

                if not parsed_chunks:
                    return self._failure(
                        ErrorCodes.PARSER_INVALID_INPUT, "No slides found in PPTX."
                    )
                return parsed_chunks
        except zipfile.BadZipFile as exc:
            return self._failure(
                ErrorCodes.PARSER_CORRUPTED_FILE, f"Failed to parse PPTX: {exc}"
            )

    def _parse_spreadsheet(
        self, content: str | bytes, _metadata: dict[str, Any], file_ext: str
    ) -> list[ParsedChunk] | AgentFailure:
        if not isinstance(content, bytes | bytearray):
            return self._failure(
                ErrorCodes.PARSER_INVALID_INPUT,
                "Spreadsheet content must be bytes.",
            )

        parsed_chunks: list[ParsedChunk] = []
        chunk_index = 0
        try:
            if file_ext == ".csv":
                df = pd.read_csv(io.BytesIO(content))
                markdown = self._table_to_markdown(df)
                parsed_chunks.append(
                    ParsedChunk(
                        chunk_id=str(uuid.uuid4()),
                        content=markdown,
                        chunk_index=chunk_index,
                        layout_type="table",
                        bbox=None,
                    )
                )
                return parsed_chunks

            excel_file = pd.ExcelFile(io.BytesIO(content))
            for sheet_name in excel_file.sheet_names:
                df = excel_file.parse(sheet_name)
                markdown = self._table_to_markdown(df)
                parsed_chunks.append(
                    ParsedChunk(
                        chunk_id=str(uuid.uuid4()),
                        content=markdown,
                        chunk_index=chunk_index,
                        layout_type="table",
                        bbox=None,
                    )
                )
                chunk_index += 1
            return parsed_chunks
        except Exception as exc:
            return self._failure(
                ErrorCodes.PARSER_INVALID_INPUT,
                f"Failed to parse spreadsheet: {exc}",
            )

    def _elements_to_chunks(self, elements: list[Any]) -> list[ParsedChunk]:
        parsed_chunks: list[ParsedChunk] = []
        chunk_index = 0

        for element in elements:
            text = getattr(element, "text", "") or ""
            text = text.strip()
            layout_type: LayoutType = self._map_layout_type(element)
            metadata = getattr(element, "metadata", None)
            page_number = getattr(metadata, "page_number", None) if metadata else None
            bbox = None

            if layout_type == "table":
                table_markdown = self._table_from_element(element)
                table_text = table_markdown.strip() if table_markdown else text
                if not table_text:
                    continue
                parsed_chunks.append(
                    ParsedChunk(
                        chunk_id=str(uuid.uuid4()),
                        content=table_text,
                        chunk_index=chunk_index,
                        layout_type="table",
                        page_number=page_number,
                        bbox=bbox,
                    )
                )
                chunk_index += 1
                continue

            if layout_type == "image":
                image_text = text or "Image"
                parsed_chunks.append(
                    ParsedChunk(
                        chunk_id=str(uuid.uuid4()),
                        content=image_text,
                        chunk_index=chunk_index,
                        layout_type="image",
                        page_number=page_number,
                        bbox=bbox,
                    )
                )
                chunk_index += 1
                continue

            if not text:
                continue

            for chunk in self._split_by_markdown_tables(text):
                chunk_text = chunk["text"].strip()
                if not chunk_text:
                    continue
                chunk_layout: LayoutType = (
                    "table" if chunk["type"] == "table" else layout_type
                )
                parsed_chunks.append(
                    ParsedChunk(
                        chunk_id=str(uuid.uuid4()),
                        content=chunk_text,
                        chunk_index=chunk_index,
                        layout_type=chunk_layout,
                        page_number=page_number,
                        bbox=bbox,
                    )
                )
                chunk_index += 1

        return parsed_chunks

    def _table_from_element(self, element: Any) -> str:
        metadata = getattr(element, "metadata", None)
        html = None
        if metadata is not None:
            html = getattr(metadata, "text_as_html", None) or getattr(
                metadata, "table_as_html", None
            )
        if html:
            markdown = self._table_from_html(html)
            if markdown:
                return markdown
        return getattr(element, "text", "") or ""

    def _table_from_html(self, html: str) -> str | None:
        try:
            dfs = pd.read_html(html)
        except Exception:
            return None
        if not dfs:
            return None
        return "\n\n".join(self._table_to_markdown(df) for df in dfs)

    def _table_to_markdown(self, df: pd.DataFrame) -> str:
        df = df.fillna("")
        headers = [str(col).strip() for col in df.columns]
        if not headers:
            return ""
        header_line = "| " + " | ".join(headers) + " |"
        separator_line = "| " + " | ".join("---" for _ in headers) + " |"
        rows = []
        for row in df.itertuples(index=False):
            row_text = [str(cell).strip() for cell in row]
            rows.append("| " + " | ".join(row_text) + " |")
        return "\n".join([header_line, separator_line, *rows])

    def _map_layout_type(self, element: Any) -> LayoutType:
        category = str(getattr(element, "category", "")).lower()
        if category in {"title", "header"}:
            return "header"
        if category in {"table"}:
            return "table"
        if category in {"image"}:
            return "image"
        if category in {"listitem", "list"}:
            return "list"
        return "text"

    def _extract_text_from_xml(self, xml_bytes: bytes) -> str:
        try:
            root = ElementTree.fromstring(xml_bytes)
        except ElementTree.ParseError:
            return ""
        texts: list[str] = []
        for node in root.iter():
            tag = node.tag
            if (tag.endswith("}t") or tag == "t") and node.text:
                cleaned = node.text.strip()
                if cleaned:
                    texts.append(cleaned)
        return "\n".join(texts)

    def _slide_sort_key(self, path: str) -> int:
        number = self._slide_number_from_path(path)
        return number or 0

    def _slide_number_from_path(self, path: str) -> int | None:
        stem = Path(path).stem
        if stem.startswith("slide"):
            number_part = stem.replace("slide", "")
            if number_part.isdigit():
                return int(number_part)
        return None

    def _is_ocr_error(self, exc: Exception) -> bool:
        name = exc.__class__.__name__.lower()
        message = str(exc).lower()
        return "tesseract" in message or "ocr" in message or "tesseract" in name

    def _is_corrupted_pdf_error(self, exc: Exception) -> bool:
        name = exc.__class__.__name__
        message = str(exc).lower()
        return name in {"PDFSyntaxError", "PdfReadError"} or "corrupt" in message

    def _failure(self, error_code: str, message: str) -> AgentFailure:
        return AgentFailure(
            agent_id="parser",
            error_code=error_code,
            message=message,
            recoverable=False,
        )
